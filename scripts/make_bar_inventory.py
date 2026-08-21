"""Build a print-ready Word inventory sheet for the bar (alcohol supplied by the couple)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "print"
DOCX_PATH = OUT_DIR / "Bar-inventaire-alcool.docx"

INK = RGBColor(0x2F, 0x2A, 0x27)
MUTED = RGBColor(0x7C, 0x71, 0x6A)
PAPER = RGBColor(0xFF, 0xFF, 0xFF)
CHALK = RGBColor(0xD2, 0xC9, 0xC2)

INK_HEX = "2F2A27"
FAINT_HEX = "E2DCD6"
RULE_HEX = "B6ABA2"
WRITE_HEX = "FBF9F7"
GROUP_HEX = "F4F0EC"

DISPLAY = "Georgia"
LABEL = "Calibri"

# (group label FR, group label EN, group colour, [(product FR, product EN or None, quantity)])
GROUPS = [
    (
        "Spiritueux",
        "Spirits",
        "8A5A1F",
        [("Whisky / Scotch", None, 5), ("Rhum", "Rum", 2), ("Vodka", None, 2)],
    ),
    (
        "Apéritifs & liqueurs",
        "Aperitifs & liqueurs",
        "7B4A57",
        [("Martini", None, 3), ("Baileys", None, 2)],
    ),
    (
        "Bulles",
        "Sparkling",
        "8A7524",
        [("Champagne", None, 2), ("Vin mousseux", "Sparkling wine", 6)],
    ),
    (
        "Vins tranquilles",
        "Still wines",
        "7A2130",
        [("Vin rouge", "Red wine", 75), ("Vin blanc", "White wine", 26)],
    ),
    (
        "Bières",
        "Beer",
        "5F6B37",
        [("Leffe", None, 300), ("Stella Artois", None, 56)],
    ),
]

# (French header, English header)
COLUMNS = [
    ("Produit", "Product"),
    ("Fournies", "Supplied"),
    ("Reçues au bar", "Received at bar"),
    ("Servies", "Served"),
    ("Restantes", "Remaining"),
]
COL_WIDTHS = [Inches(2.60), Inches(1.15), Inches(1.15), Inches(1.15), Inches(1.15)]


def set_run_font(
    run,
    name: str,
    size_pt: float,
    color: RGBColor,
    *,
    bold: bool = False,
    spacing_pt: float | None = None,
    caps: bool = False,
) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.font.east_asia = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if spacing_pt:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(int(spacing_pt * 20)))
        rpr.append(spacing)
    if caps:
        rpr.append(OxmlElement("w:caps"))


def shade(cell, hex_fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_borders(cell, edges: dict[str, tuple[str, int]]) -> None:
    """edges maps 'top'/'bottom'/'left'/'right' to (colour hex, size in eighths of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if edge in edges:
            color, size = edges[edge]
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(size))
            element.set(qn("w:color"), color)
        else:
            element.set(qn("w:val"), "nil")


def set_cell_margins(table, top: int, right: int, bottom: int, left: int) -> None:
    """Margins in twentieths of a point."""
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def write_cell(
    cell,
    text: str,
    *,
    font: str = DISPLAY,
    size: float = 11,
    color: RGBColor = INK,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    spacing_pt: float | None = None,
    caps: bool = False,
    space_before: float = 3,
    space_after: float = 3,
):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, color, bold=bold, spacing_pt=spacing_pt, caps=caps)
    return paragraph


def append_run(
    paragraph,
    text: str,
    *,
    font: str = LABEL,
    size: float = 8,
    color: RGBColor = MUTED,
    spacing_pt: float | None = None,
    caps: bool = False,
) -> None:
    """Add the English half of a bilingual label to an existing paragraph."""
    run = paragraph.add_run(text)
    set_run_font(run, font, size, color, spacing_pt=spacing_pt, caps=caps)


def add_cell_line(
    cell,
    text: str,
    *,
    font: str = LABEL,
    size: float = 7,
    color: RGBColor = MUTED,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    spacing_pt: float | None = None,
    caps: bool = False,
    space_after: float = 3,
) -> None:
    """Add a second line (the English one) below a cell's first line."""
    paragraph = cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, color, spacing_pt=spacing_pt, caps=caps)


def add_paragraph(
    document: Document,
    text: str,
    *,
    font: str,
    size: float,
    color: RGBColor,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    space_before: float = 0,
    space_after: float = 6,
    spacing_pt: float | None = None,
    caps: bool = False,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, color, bold=bold, spacing_pt=spacing_pt, caps=caps)
    return paragraph


def add_summary(document: Document, totals: list[tuple[str, str, int]], grand_total: int) -> None:
    table = document.add_table(rows=2, cols=len(totals) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table, 60, 100, 60, 100)

    cells = [(("Total bouteilles", "Total bottles", grand_total), True)]
    cells += [((label_fr, label_en, value), False) for label_fr, label_en, value in totals]
    width = Inches(7.20 / (len(totals) + 1))

    for index, ((label_fr, label_en, value), is_total) in enumerate(cells):
        label_cell = table.cell(0, index)
        value_cell = table.cell(1, index)
        for cell in (label_cell, value_cell):
            cell.width = width
            shade(cell, INK_HEX if is_total else "FFFFFF")
        set_cell_borders(
            label_cell,
            {"top": (RULE_HEX, 8), "left": (FAINT_HEX, 6), "right": (FAINT_HEX, 6)},
        )
        set_cell_borders(
            value_cell,
            {"bottom": (RULE_HEX, 8), "left": (FAINT_HEX, 6), "right": (FAINT_HEX, 6)},
        )
        write_cell(
            label_cell,
            label_fr,
            font=LABEL,
            size=7.5,
            color=CHALK if is_total else MUTED,
            bold=True,
            spacing_pt=1.1,
            caps=True,
            space_before=4,
            space_after=0,
        )
        add_cell_line(
            label_cell,
            label_en,
            size=6.5,
            color=CHALK if is_total else MUTED,
            spacing_pt=0.9,
            caps=True,
            space_after=0,
        )
        write_cell(
            value_cell,
            str(value),
            font=DISPLAY,
            size=20,
            color=PAPER if is_total else INK,
            space_before=0,
            space_after=4,
        )


def add_inventory_table(document: Document, grand_total: int) -> None:
    row_count = 1 + sum(1 + len(items) for *_, items in GROUPS) + 1
    table = document.add_table(rows=row_count, cols=len(COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table, 70, 110, 70, 110)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = COL_WIDTHS[index]

    # Header
    header = table.rows[0]
    for index, (title_fr, title_en) in enumerate(COLUMNS):
        cell = header.cells[index]
        align = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        shade(cell, INK_HEX)
        set_cell_borders(cell, {})
        write_cell(
            cell,
            title_fr,
            font=LABEL,
            size=8,
            color=PAPER,
            bold=True,
            spacing_pt=1.0,
            caps=True,
            align=align,
            space_before=4,
            space_after=0,
        )
        add_cell_line(
            cell,
            title_en,
            size=6.5,
            color=CHALK,
            spacing_pt=0.8,
            caps=True,
            align=align,
            space_after=4,
        )

    row_index = 1
    for label_fr, label_en, color_hex, items in GROUPS:
        subtotal = sum(quantity for *_, quantity in items)

        group_row = table.rows[row_index]
        name_cell = group_row.cells[0]
        count_cell = group_row.cells[1].merge(group_row.cells[4])
        for cell in (name_cell, count_cell):
            shade(cell, GROUP_HEX)
            set_cell_borders(cell, {"bottom": (INK_HEX, 6)})
        set_cell_borders(name_cell, {"bottom": (INK_HEX, 6), "left": (color_hex, 24)})
        group_paragraph = write_cell(
            name_cell,
            label_fr,
            font=LABEL,
            size=8.5,
            color=RGBColor.from_string(color_hex),
            bold=True,
            spacing_pt=1.2,
            caps=True,
            space_before=4,
            space_after=4,
        )
        append_run(
            group_paragraph,
            f"  ·  {label_en}",
            size=7.5,
            color=MUTED,
            spacing_pt=0.9,
            caps=True,
        )
        write_cell(
            count_cell,
            f"{subtotal} bouteilles / bottles",
            font=LABEL,
            size=8.5,
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            space_before=4,
            space_after=4,
        )
        row_index += 1

        for product_fr, product_en, quantity in items:
            row = table.rows[row_index]
            item_cell = row.cells[0]
            set_cell_borders(
                item_cell,
                {"bottom": (FAINT_HEX, 6), "left": (color_hex, 24)},
            )
            item_paragraph = write_cell(item_cell, product_fr, size=12)
            if product_en:
                append_run(item_paragraph, f"   {product_en}", size=8, color=MUTED)

            quantity_cell = row.cells[1]
            set_cell_borders(quantity_cell, {"bottom": (FAINT_HEX, 6)})
            write_cell(
                quantity_cell,
                str(quantity),
                size=13,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
            )

            for column in (2, 3, 4):
                blank = row.cells[column]
                shade(blank, WRITE_HEX)
                set_cell_borders(
                    blank,
                    {"bottom": (FAINT_HEX, 6), "left": (FAINT_HEX, 6)},
                )
                write_cell(blank, "", size=13)
            row_index += 1

    total_row = table.rows[row_index]
    for index, cell in enumerate(total_row.cells):
        shade(cell, "FFFFFF")
        set_cell_borders(cell, {"top": (INK_HEX, 12)})
    total_paragraph = write_cell(
        total_row.cells[0],
        "Total général",
        font=LABEL,
        size=9,
        bold=True,
        spacing_pt=1.2,
        caps=True,
        space_before=6,
        space_after=4,
    )
    append_run(
        total_paragraph,
        "  ·  Grand total",
        size=8,
        color=MUTED,
        spacing_pt=1.0,
        caps=True,
    )
    write_cell(
        total_row.cells[1],
        str(grand_total),
        size=15,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=5,
        space_after=4,
    )


def add_signature_block(document: Document) -> None:
    captions = [
        ("Remis au bar par · heure", "Delivered to bar by · time"),
        ("Reçu par (barman) · heure", "Received by (bartender) · time"),
        ("Décompte final · heure", "Final count · time"),
    ]
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(table, 40, 140, 40, 0)

    for index, (caption_fr, caption_en) in enumerate(captions):
        line_cell = table.cell(0, index)
        caption_cell = table.cell(1, index)
        line_cell.width = Inches(2.4)
        caption_cell.width = Inches(2.4)
        set_cell_borders(line_cell, {"bottom": (RULE_HEX, 8)})
        set_cell_borders(caption_cell, {})
        write_cell(line_cell, "", size=12, space_before=14, space_after=2)
        write_cell(
            caption_cell,
            caption_fr,
            font=LABEL,
            size=7.5,
            color=MUTED,
            spacing_pt=0.9,
            caps=True,
            space_before=2,
            space_after=0,
        )
        add_cell_line(
            caption_cell,
            caption_en,
            size=6.5,
            color=MUTED,
            spacing_pt=0.7,
            caps=True,
            space_after=0,
        )


def build_document(dest: Path) -> int:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    grand_total = sum(quantity for *_, items in GROUPS for *_, quantity in items)

    add_paragraph(
        document,
        "Georges & Christella · Bar du mariage · Wedding bar",
        font=LABEL,
        size=8.5,
        color=MUTED,
        spacing_pt=1.6,
        caps=True,
        space_after=2,
    )
    title = add_paragraph(
        document,
        "Inventaire des bouteilles",
        font=DISPLAY,
        size=25,
        color=INK,
        space_after=3,
    )
    append_run(title, "   Bottle inventory", font=DISPLAY, size=15, color=MUTED)
    add_paragraph(
        document,
        "Samedi 15 août 2026 · Centre des congrès et banquets Renaissance · Alcool fourni par les mariés",
        font=LABEL,
        size=9,
        color=MUTED,
        space_after=1,
    )
    add_paragraph(
        document,
        "Saturday, August 15, 2026 · All alcohol supplied by the bride and groom",
        font=LABEL,
        size=8,
        color=MUTED,
        space_after=9,
    )

    totals = [
        ("Spiritueux, apéritifs & liqueurs", "Spirits, aperitifs & liqueurs", 14),
        ("Bulles", "Sparkling", 8),
        ("Vins tranquilles", "Still wines", 101),
        ("Bières", "Beer", 356),
    ]
    add_summary(document, totals, grand_total)

    add_paragraph(document, "", font=LABEL, size=5, color=MUTED, space_after=0)
    add_inventory_table(document, grand_total)
    add_paragraph(document, "", font=LABEL, size=7, color=MUTED, space_after=0)
    add_signature_block(document)

    add_paragraph(
        document,
        "Toutes les bouteilles sont fournies par les mariés. Merci de compter à la réception, "
        "puis de nouveau en fin de soirée, et de remettre les bouteilles non ouvertes à la famille.",
        font=LABEL,
        size=8,
        color=MUTED,
        space_before=13,
        space_after=1,
    )
    add_paragraph(
        document,
        "All bottles are supplied by the bride and groom. Please count them on delivery, count again "
        "at the end of the night, and return every unopened bottle to the family.",
        font=LABEL,
        size=8,
        color=MUTED,
        space_after=0,
    )

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(dest)
    return grand_total


def main() -> None:
    total = build_document(DOCX_PATH)
    print(f"Wrote {DOCX_PATH} · {total} bouteilles")


if __name__ == "__main__":
    main()
