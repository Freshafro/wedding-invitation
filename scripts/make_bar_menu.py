"""Build the guest-facing open-bar display card the bartenders stand on the bar.

Companion to make_bar_inventory.py: same drinks, but no counts and no tracking
columns — this one is read across a bar, not written on.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from make_bar_inventory import DISPLAY, LABEL, add_paragraph, append_run

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "print" / "Bar-open-bar-affiche.docx"

INK = RGBColor(0x2F, 0x2A, 0x27)
MUTED = RGBColor(0x7C, 0x71, 0x6A)

# (category FR, category EN, colour, [(drink FR, drink EN or None)])
MENU = [
    (
        "Spiritueux",
        "Spirits",
        "8A5A1F",
        [("Whisky & Scotch", None), ("Rhum", "Rum"), ("Vodka", None)],
    ),
    (
        "Apéritifs & liqueurs",
        "Aperitifs & liqueurs",
        "7B4A57",
        [("Martini", None), ("Baileys", None)],
    ),
    (
        "Bulles",
        "Sparkling",
        "8A7524",
        [("Champagne", None), ("Vin mousseux", "Sparkling wine")],
    ),
    (
        "Vins",
        "Wines",
        "7A2130",
        [("Vin rouge", "Red wine"), ("Vin blanc", "White wine")],
    ),
    (
        "Bières",
        "Beers",
        "5F6B37",
        [("Leffe", None), ("Stella Artois", None)],
    ),
]


def centered(document: Document, text: str, **kwargs):
    return add_paragraph(document, text, align=WD_ALIGN_PARAGRAPH.CENTER, **kwargs)


def add_divider(document: Document, color_hex: str) -> None:
    """A short letterspaced rule in the category colour, used between sections."""
    paragraph = centered(
        document,
        "———",
        font=DISPLAY,
        size=9,
        color=RGBColor.from_string(color_hex),
        space_before=10,
        space_after=2,
        spacing_pt=2.0,
    )
    paragraph.paragraph_format.line_spacing = Pt(10)


def build_document(dest: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    centered(
        document,
        "Georges & Christella",
        font=LABEL,
        size=10,
        color=MUTED,
        spacing_pt=2.4,
        caps=True,
        space_after=6,
    )
    centered(document, "Bar ouvert", font=DISPLAY, size=40, color=INK, space_after=2)
    centered(document, "Open bar", font=DISPLAY, size=19, color=MUTED, space_after=14)

    for index, (category_fr, category_en, color_hex, drinks) in enumerate(MENU):
        if index:
            add_divider(document, color_hex)

        heading = centered(
            document,
            category_fr,
            font=LABEL,
            size=11,
            color=RGBColor.from_string(color_hex),
            bold=True,
            spacing_pt=2.2,
            caps=True,
            space_before=6 if index else 0,
            space_after=7,
        )
        append_run(
            heading,
            f"  ·  {category_en}",
            size=9.5,
            color=MUTED,
            spacing_pt=1.6,
            caps=True,
        )

        for drink_fr, drink_en in drinks:
            line = centered(
                document,
                drink_fr,
                font=DISPLAY,
                size=19,
                color=INK,
                space_after=4,
            )
            if drink_en:
                append_run(line, f"   {drink_en}", font=DISPLAY, size=13, color=MUTED)

    centered(
        document,
        "Tout est offert par les mariés · Everything is on the bride and groom",
        font=LABEL,
        size=9,
        color=MUTED,
        space_before=22,
        space_after=4,
    )
    centered(
        document,
        "Santé !",
        font=DISPLAY,
        size=15,
        color=RGBColor.from_string("7A2130"),
        space_after=0,
    )

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(dest)


def main() -> None:
    build_document(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
