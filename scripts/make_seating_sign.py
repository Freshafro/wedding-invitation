"""Build a print-ready Word sign with the seating QR code."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SVG_PATH = ROOT / "public" / "seating-qr.svg"
OUT_DIR = ROOT / "print"
PNG_PATH = OUT_DIR / "seating-qr.png"
DOCX_PATH = OUT_DIR / "Find-your-table.docx"

INK = RGBColor(0x2C, 0x24, 0x28)
MUTED = RGBColor(0x4A, 0x41, 0x46)


def parse_qr_modules(svg_text: str) -> list[list[bool]]:
    viewbox = re.search(r'viewBox="0 0 (\d+) (\d+)"', svg_text)
    if not viewbox:
        raise ValueError("Could not read QR viewBox")
    width, height = int(viewbox.group(1)), int(viewbox.group(2))
    grid = [[False] * width for _ in range(height)]

    path_match = re.search(r'<path stroke="#000000" d="([^"]+)"', svg_text)
    if not path_match:
        raise ValueError("Could not find QR path data")

    tokens = re.findall(r"[Mmh]|-?\d+(?:\.\d+)?", path_match.group(1))
    x = y = 0.0
    i = 0
    while i < len(tokens):
        token = tokens[i]
        if token == "M":
            x = float(tokens[i + 1])
            y = float(tokens[i + 2])
            i += 3
            continue
        if token == "m":
            x += float(tokens[i + 1])
            y += float(tokens[i + 2])
            i += 3
            continue
        if token == "h":
            length = float(tokens[i + 1])
            row = int(y)
            start = int(round(x))
            count = int(round(length))
            for col in range(start, start + count):
                if 0 <= row < height and 0 <= col < width:
                    grid[row][col] = True
            x += length
            i += 2
            continue
        raise ValueError(f"Unexpected SVG path token: {token}")

    return grid


def save_qr_png(grid: list[list[bool]], dest: Path, module_px: int = 24) -> None:
    height = len(grid)
    width = len(grid[0])
    image = Image.new("RGB", (width * module_px, height * module_px), "white")
    pixels = image.load()
    for row, line in enumerate(grid):
        for col, filled in enumerate(line):
            if not filled:
                continue
            for dy in range(module_px):
                for dx in range(module_px):
                    pixels[col * module_px + dx, row * module_px + dy] = (0, 0, 0)
    dest.parent.mkdir(parents=True, exist_ok=True)
    image.save(dest, "PNG")


def set_run_font(run, name: str, size_pt: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    run.font.east_asia = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_centered_paragraph(
    document: Document,
    text: str,
    *,
    font: str,
    size: float,
    color: RGBColor,
    bold: bool = False,
    space_before: float = 0,
    space_after: float = 8,
    line_pt: float | None = None,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if line_pt:
        paragraph.paragraph_format.line_spacing = Pt(line_pt)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, color, bold=bold)


def build_document(png_path: Path, dest: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_centered_paragraph(
        document,
        "Georges & Christella",
        font="Georgia",
        size=20,
        color=INK,
        space_after=6,
    )
    add_centered_paragraph(
        document,
        "15 août 2026  ·  August 15, 2026",
        font="Calibri",
        size=12,
        color=MUTED,
        space_after=22,
    )
    add_centered_paragraph(
        document,
        "Find your table",
        font="Georgia",
        size=32,
        color=INK,
        bold=True,
        space_after=2,
    )
    add_centered_paragraph(
        document,
        "Trouvez votre table",
        font="Georgia",
        size=26,
        color=INK,
        space_after=18,
    )
    add_centered_paragraph(
        document,
        "Scan this code, enter your name, and we will show you where you are seated.",
        font="Calibri",
        size=14,
        color=INK,
        space_after=6,
        line_pt=20,
    )
    add_centered_paragraph(
        document,
        "Scannez ce code, saisissez votre nom, et nous vous indiquerons votre place.",
        font="Calibri",
        size=14,
        color=INK,
        space_after=16,
        line_pt=20,
    )

    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_before = Pt(6)
    picture.paragraph_format.space_after = Pt(18)
    run = picture.add_run()
    run.add_picture(str(png_path), width=Inches(3.6))

    add_centered_paragraph(
        document,
        "Need help? Please see a member of the wedding party.",
        font="Calibri",
        size=12,
        color=MUTED,
        space_after=4,
    )
    add_centered_paragraph(
        document,
        "Besoin d’aide ? Veuillez voir un membre du cortège.",
        font="Calibri",
        size=12,
        color=MUTED,
        space_after=0,
    )

    dest.parent.mkdir(parents=True, exist_ok=True)
    document.save(dest)


def main() -> None:
    svg_text = SVG_PATH.read_text(encoding="utf-8")
    grid = parse_qr_modules(svg_text)
    save_qr_png(grid, PNG_PATH)
    build_document(PNG_PATH, DOCX_PATH)
    print(f"Wrote {PNG_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
